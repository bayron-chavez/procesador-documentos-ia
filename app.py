import streamlit as st
import json
import os
from pydantic import BaseModel, Field
import io
import docx

# Librerías de Google
from google import genai
from google.genai import types
from google.genai.errors import APIError

# Librería de OpenAI
from openai import OpenAI 

# ==========================================
# 1. SISTEMA DE AUTENTICACIÓN
# ==========================================
if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = False

if not st.session_state["autenticado"]:
    st.set_page_config(page_title="Acceso Restringido", page_icon="🔒", layout="centered")
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.subheader("🔒 Acceso al Procesador de Documentos")
    st.write("Esta aplicación contiene herramientas corporativas de auditoría. Por favor, identifícate.")
    
    usuario = st.text_input("Usuario", placeholder="Introduce tu usuario")
    contrasena = st.text_input("Contraseña", type="password", placeholder="Introduce tu contraseña")
    
    if st.button("Ingresar al Sistema", type="primary", use_container_width=True):
        if usuario == "admin" and contrasena == "admin2026": # Credenciales de ejemplo, reemplaza con un sistema real en producción
            st.session_state["autenticado"] = True
            st.rerun()
        else:
            st.error("❌ Credenciales incorrectas. Acceso denegado.")
            
    st.stop() # Detiene por completo el script si el usuario no se ha autenticado


# ==========================================
# 2. CONFIGURACIÓN DE LA INTERFAZ PRINCIPAL
# ==========================================
st.set_page_config(
    page_title="Procesador de documentos", 
    page_icon="📄", 
    layout="wide"
)

st.title("Procesador de Documentos")
st.markdown("Sube un documento y extrae la información clave devolviendo un JSON estructurado utilizando los motores de IA más avanzados del mercado.")

# Configuración del menú lateral pidiendo la API Key
with st.sidebar:
    st.header("⚙️ Motor de IA")
    
    # DICCIONARIO DE ENRUTAMIENTO: La forma más limpia y profesional de escalar modelos
    modelos_disponibles = {
        "Gemini 2.5 Flash (Rápido y eficiente)": {"proveedor": "Google Gemini", "id": "gemini-2.5-flash"},
        "Gemini 2.5 Pro (Más potente Google)": {"proveedor": "Google Gemini", "id": "gemini-2.5-pro"},
        "OpenAI GPT-4o-mini (Rápido y económico)": {"proveedor": "OpenAI (GPT)", "id": "gpt-4o-mini"},
        "OpenAI GPT-4o (Insignia omnimodal)": {"proveedor": "OpenAI (GPT)", "id": "gpt-4o"},
        "OpenAI o3-mini (Razonamiento profundo)": {"proveedor": "OpenAI (GPT)", "id": "o3-mini"}
    }
    
    modelo_visual = st.selectbox(
        "Selecciona la Inteligencia Artificial",
        options=list(modelos_disponibles.keys())
    )
    
    # Extraemos automáticamente el proveedor y el ID técnico desde el diccionario
    proveedor_ia = modelos_disponibles[modelo_visual]["proveedor"]
    modelo_seleccionado = modelos_disponibles[modelo_visual]["id"]
    
    st.markdown("---")
    st.header("🔑 Credenciales")
    
    # La interfaz pide la clave correcta según el proveedor extraído
    if proveedor_ia == "Google Gemini":
        api_key_input = st.text_input(
            "API Key de Google GenAI", 
            type="password",
            help="Introduce tu API Key de Google."
        )
    elif proveedor_ia == "OpenAI (GPT)":
        api_key_input = st.text_input(
            "API Key de OpenAI", 
            type="password",
            help="Introduce tu API Key de OpenAI."
        )
        
    st.markdown(f"**ID Activo:** `{modelo_seleccionado}`")
        
    st.markdown("---")
    if st.button("Cerrar Sesión", type="secondary"):
        st.session_state["autenticado"] = False
        st.rerun()


# 3. DEFINICIÓN DEL ESQUEMA DE DATOS (MOLDE DE NEGOCIO)
class DocumentAnalysis(BaseModel):
    entidad_emisora: str = Field(..., description="Nombre de la entidad emisora del documento")
    identificador_unico: str = Field(..., description="Identificador único del documento")
    fecha_emision: str = Field(..., description="Fecha de emisión del documento")
    monto_total: float = Field(..., description="Monto total asociado al documento")
    resumen_ejecutivo: str = Field(..., description="Resumen conciso del contenido del documento")
    puntos_clave: list[str] = Field(..., description="Lista de puntos clave extraídos del documento")


# 4. COMPONENTE DE SUBIDA DE ARCHIVOS
uploaded_file = st.file_uploader(
    "Sube tu documento aquí", 
    type=["pdf", "txt", "docx", "doc"]
)

if uploaded_file is not None:
    columna_doc, columna_ia = st.columns([1, 1])
    file_content = uploaded_file.read()

    # COLUMNA IZQUIERDA: VISTA PREVIA
    with columna_doc:
        st.subheader("Documento Subido 🔼")
        st.write(f"**Nombre:** {uploaded_file.name}") 

        if uploaded_file.type == "application/pdf":
            st.info("Visualización directa no disponible. Gemini analizará de forma nativa la estructura completa del PDF.")
        elif uploaded_file.type in ["text/plain"]:
            st.text(file_content.decode("utf-8", errors="ignore"))
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            st.info("Vista previa no disponible para archivos Word.")
        else:
            st.write("Tipo de archivo no soportado para vista previa.")

    # COLUMNA DERECHA: ANÁLISIS E INTEGRACIÓN CON IA
    with columna_ia:
        st.subheader("Análisis del Documento 🤖")

        if not api_key_input:
            st.warning(f"⚠️ Por favor, introduce tu API Key de {proveedor_ia} en la barra lateral para habilitar el procesamiento.")
        else:
            if st.button(f"Procesar con {proveedor_ia}", type="primary", use_container_width=True):
                with st.spinner(f"Procesando el documento con {proveedor_ia}, por favor espera..."):
                    try:
                        # ==========================================
                        # RUTA 1: GOOGLE GEMINI
                        # ==========================================
                        if proveedor_ia == "Google Gemini":
                            client = genai.Client(api_key=api_key_input)
                            
                            # Validamos el tipo de archivo
                            if uploaded_file.type == "application/pdf":
                                contenido_input = [
                                    types.Part.from_bytes(data=file_content, mime_type="application/pdf"),
                                    "Extrae minuciosamente la información clave de este documento siguiendo el esquema JSON requerido."
                                ]
                            elif uploaded_file.name.endswith(".docx"):
                                # MAGIA PARA WORD: Extraemos el texto limpio
                                doc = docx.Document(io.BytesIO(file_content))
                                texto_word = "\n".join([para.text for para in doc.paragraphs])
                                contenido_input = [f"Extrae la información clave de este documento:\n\n{texto_word}"]
                            else:
                                contenido_input = [f"Extrae la información clave de este documento:\n\n{file_content.decode('utf-8', errors='ignore')}"]

                            response = client.models.generate_content(
                                model=modelo_seleccionado,
                                contents=contenido_input,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json",
                                    response_schema=DocumentAnalysis,
                                    temperature=0.1,
                                    system_instruction="Eres un sistema automatizado de auditoría y extracción de datos. Mapea los documentos al formato JSON solicitado. Si no existe un dato, déjalo vacío o en 0.0."
                                )
                            )
                            datos_finales = json.loads(response.text)

                        # ==========================================
                        # RUTA 2: OPENAI (GPT)
                        # ==========================================
                        elif proveedor_ia == "OpenAI (GPT)":
                            client_openai = OpenAI(api_key=api_key_input)
                            
                            if uploaded_file.type == "application/pdf":
                                st.error("⚠️ Para procesar PDFs con OpenAI requieres extraer el texto primero.")
                                st.stop()
                            elif uploaded_file.name.endswith(".docx"):
                                # MAGIA PARA WORD EN OPENAI
                                doc = docx.Document(io.BytesIO(file_content))
                                texto_archivo = "\n".join([para.text for para in doc.paragraphs])
                            else:
                                texto_archivo = file_content.decode("utf-8", errors="ignore")
                            
                            completion = client_openai.beta.chat.completions.parse(
                                model=modelo_seleccionado,
                                messages=[
                                    {"role": "system", "content": "Eres un sistema automatizado de auditoría y extracción de datos. Mapea los documentos al formato JSON solicitado. Si no existe un dato, déjalo vacío o en 0.0."},
                                    {"role": "user", "content": f"Extrae la información clave de este documento:\n\n{texto_archivo}"}
                                ],
                                response_format=DocumentAnalysis,
                                temperature=0.1
                            )
                            datos_finales = json.loads(completion.choices[0].message.content)

                        # ==========================================
                        # RENDERIZADO VISUAL COMÚN
                        # ==========================================
                        st.success(f"¡Documento procesado exitosamente usando {proveedor_ia}!")
                        st.markdown("---")

                        m1, m2, m3 = st.columns(3)
                        m1.metric("Emisor", datos_finales.get("entidad_emisora", "No detectado"))
                        m2.metric("Fecha", datos_finales.get("fecha_emision", "No detectada"))
                        m3.metric("Monto Total", f"${datos_finales.get('monto_total', 0.0):,.2f}")

                        st.markdown(f"**Identificador Único:** `{datos_finales.get('identificador_unico', 'No detectado')}`")
                        
                        st.markdown("### 📝 Resumen Ejecutivo")
                        st.write(datos_finales.get("resumen_ejecutivo"))

                        st.markdown("### 📌 Puntos Clave")
                        for punto in datos_finales.get("puntos_clave", []):
                            st.markdown(f"- {punto}")

                        with st.expander("Ver JSON Completo listo para API o Base de Datos"):
                            st.json(datos_finales)

                    except Exception as e:
                        st.error(f"Ocurrió un error inesperado durante el análisis: {str(e)}")